from __future__ import annotations

import sys
import re
from pathlib import Path
from typing import List

try:
    from docx import Document  # type: ignore
except Exception as exc:  # pragma: no cover
    print("ERROR: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    raise


def sanitize_filename(name: str) -> str:
    # Replace forbidden/special characters with dashes
    s = re.sub(r"[\\/:*?\"<>|]", "-", name)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def para_to_md(p) -> str:
    text = p.text or ""
    style = (getattr(p.style, "name", "") or "").lower()
    if style.startswith("heading "):
        try:
            level = int(style.split()[1])
        except Exception:
            level = 1
        level = max(1, min(level, 6))
        return f"{'#'*level} {text}\n\n"
    if "list bullet" in style or "list paragraph" in style:
        return f"- {text}\n"
    if "list number" in style:
        # Keep simple: mark as bullet to avoid renumbering complexity
        return f"- {text}\n"
    # Default paragraph
    return f"{text}\n\n"


def table_to_md(tbl) -> str:
    rows = []
    for r in tbl.rows:
        cells = [c.text.replace("\n", " ").strip() for c in r.cells]
        rows.append(cells)
    if not rows:
        return "\n"
    # Build a simple markdown table; if only one row, no header line
    out: List[str] = []
    out.append("| " + " | ".join(rows[0]) + " |\n")
    if len(rows) > 1:
        out.append("| " + " | ".join(["---"] * len(rows[0])) + " |\n")
        for row in rows[1:]:
            out.append("| " + " | ".join(row) + " |\n")
    return "".join(out) + "\n\n"


def convert_docx_to_markdown(src: Path, dst: Path) -> None:
    doc = Document(str(src))
    parts: List[str] = []
    # Document body: paragraphs and tables in order
    # python-docx doesn't directly iterate mixed flow; approximate by collecting paragraphs then tables in sequence
    # Here we traverse blocks via element tree
    from docx.oxml.text.paragraph import CT_P  # type: ignore
    from docx.oxml.table import CT_Tbl  # type: ignore

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            p = child
            # Wrap CT_P back to paragraph API
            para = doc.paragraphs[len(parts)] if len(parts) < len(doc.paragraphs) else None
            # Fallback: find matching paragraph by text
            if para is None:
                text = "".join([r.text for r in p.iter() if hasattr(r, 'text')])
                parts.append(text + "\n\n")
            else:
                parts.append(para_to_md(para))
        elif isinstance(child, CT_Tbl):
            # Find the next table by index
            idx = sum(1 for s in parts if s == "__TABLE_PLACEHOLDER__")
            parts.append("__TABLE_PLACEHOLDER__")
        else:
            # Unknown element; skip
            continue

    # Replace placeholders with actual tables in order
    tbl_iter = iter(doc.tables)
    for i, part in enumerate(parts):
        if part == "__TABLE_PLACEHOLDER__":
            try:
                tbl = next(tbl_iter)
                parts[i] = table_to_md(tbl)
            except StopIteration:
                parts[i] = "\n"

    # If parts was empty (structure mismatch), fall back to paragraph-first, then tables
    if not parts:
        for p in doc.paragraphs:
            parts.append(para_to_md(p))
        for tbl in doc.tables:
            parts.append(table_to_md(tbl))

    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text("".join(parts), encoding="utf-8")


def main(argv: List[str]) -> int:
    if len(argv) < 2:
        print("Usage: docx_to_markdown.py <file.docx> [more.docx ...]", file=sys.stderr)
        return 2
    for src_str in argv[1:]:
        src = Path(src_str)
        if not src.exists():
            print(f"WARN: {src} not found", file=sys.stderr)
            continue
        base = sanitize_filename(src.stem) + ".md"
        dst = Path(__file__).resolve().parents[1] / "docs" / "uploads" / base
        try:
            convert_docx_to_markdown(src, dst)
            print(f"Converted: {src} -> {dst}")
        except Exception as exc:
            print(f"ERROR converting {src}: {exc}", file=sys.stderr)
            return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
